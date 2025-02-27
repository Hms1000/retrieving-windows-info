import os
import wmi
import sys
import logging
import argparse
from docx import Document

#set up logging
def configure_logging(output_dir):
    log_path = os.path.join(output_dir, 'List_of_running_services.log')
    os.makedirs(output_dir, exist_ok=True)
    logging.info(f'Created output directory: {output_dir}')

    logging.basicConfig(
        filename = log_path,
        level = logging.INFO,
        format = '%(asctime)s - %(levelname)s - %(message)s'
    )

#create a word document
def create_word_document():
    doc = Document()
    doc.add_heading('Table of Running Services',0)
    return doc

def retrieve_services():
    try:
        conn = wmi.WMI()
        services = conn.Win32_Service()

        if not services:
            logging.warning('No Services found')
            print('No services found')
            sys.exit(1)

        logging.info(f'Retrieved {len(services)} services successfully.')
        return services

    except wmi.x_wmi as e:
        logging.error(f'Error accessing WMI services: {e}')
        print(f'Error accessing WMI services: {e}')
        sys.exit(1)
    except Exception as e:
        logging.error(f'Unexpected error occurred: {e}')
        print(f'Unexpected error occurred:{e}')
        sys.exit(1)

#create a table to store driver data
def create_table(doc, desired_attributes):
    table = doc.add_table(rows=1, cols=len(desired_attributes))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for index, attr in enumerate(desired_attributes):
        header_cells[index].text = attr
    return table

#populate table with service information
def populate_table_with_service_info(table, services, desired_attributes):
    for service in services:
        row_cells = table.add_row().cells
        for index, attr in enumerate(desired_attributes):
            try:
                value = getattr(service, attr, None)
                if value is None or (isinstance(value, str) and not value.strip()):
                    row_cells[index].text = 'N/A'
                else:
                    row_cells[index].text = str(value)

            except AttributeError as e:
                row_cells[index].text = 'Error'
                logging.warning(f'Missing attribute {attr} foa service')
            except Exception as e:
                row_cells[index].text = 'Error'
                logging.error(f'Error retrieving {attr} for service: {e}')
                print(f'Error retrieving {attr} for service: {e}')

#save document
def save_document(doc, output_path):
    try:
        doc.save(output_path)
        logging.info(f'Document saved successfully: {output_path}')
        print(f'Document saved successfully: {output_path}')
    except Exception as e:
        logging.error(f'Error saving Document: {e}')
        print(f'Error saving Document: {e}')

#main function to consolidate everything
def main():
    try:
        parser = argparse.ArgumentParser(description='Retrieve Services Information')
        parser.add_argument('--output', help='Output file', default=os.path.join(os.path.expanduser('~'), 'Documents', 'List_of_running_services.docx'))
        args = parser.parse_args()

        #makesure the output directory exists
        output_dir = os.path.dirname(args.output)

        #configure logging
        configure_logging(output_dir)

        #create document
        doc = create_word_document()

        #Desired_attributes
        # Step 5: Desired service attributes
        desired_attributes = [
            "Name", "DisplayName", "State", "StartMode", "Status", "PathName", "StartName"
        ]

        #retrieving running services
        services = retrieve_services()

        #create a table
        table = create_table(doc, desired_attributes)

        #populate the table
        populate_table_with_service_info(table, services, desired_attributes)

        #save the document
        save_document(doc, args.output)

    except Exception as e:
        logging.error(f'Unexpected error in main: {e}')
        print(f'Unexpected error in main: {e}')

if __name__ == '__main__':
    main()




