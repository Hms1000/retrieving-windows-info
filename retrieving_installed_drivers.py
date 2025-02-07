'''This is a script to retrieve installed drivers on windows machine'''
import os
import wmi
import sys
import logging
import argparse
from docx import Document

#Set up logging configuration.
def setup_logging(output_dir):
    log_path = os.path.join(output_dir, "List_of_installed_drivers.log")
    logging.basicConfig(
        filename=log_path,
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )
    logging.info(f'Logging configured. Log file: {log_path}')

#Create a Word document with a heading for the installed drivers table.
def create_word_document():
    doc = Document()
    doc.add_heading("Table of Installed Drivers", 0)
    return doc

#Retrieve installed drivers using WMI
def get_installed_drivers():
    try:
        conn = wmi.WMI()
        drivers = conn.Win32_PnPSignedDriver()

        if not drivers:
            print("No drivers found.")
            logging.warning("No drivers found.")
            sys.exit(1)

        return drivers

    except wmi.x_wmi as e:
        logging.error(f"Error running WMI process: {e}")
        print(f"Error running WMI process: {e}")
        sys.exit(1)

    except Exception as e:
        logging.error(f'Unexpected Error: {e}')
        print(f'Unexpected Error: {e}')
        sys.exit(1)


#Create a table in the Word document to store driver data
def create_table(doc, desired_attributes):

    table = doc.add_table(rows=1, cols=len(desired_attributes))
    table.style = "Table Grid"

    # Add table headers
    header_cells = table.rows[0].cells
    for index, attr in enumerate(desired_attributes):
        header_cells[index].text = attr

    return table

# Populate the Word document table with driver data
def populate_table_with_driver_data(table, drivers, desired_attributes):
    for driver in drivers:
        row_cells = table.add_row().cells
        for index, attr in enumerate(desired_attributes):
            try:
                #retrieve attribute value from the driver using getattr
                value = getattr(driver, attr, None)
                # Use getattr to safely retrieve attribute values
                if value is None or (isinstance(value, str) and not value.strip()):
                    row_cells[index].text = 'N/A'
                else:
                    row_cells[index].text = str(value)
            except Exception as e:
                row_cells[index].text = "Error"
                logging.error(f"Error retrieving {attr} for driver: {e}")

#Save the Word document to the specified path
def save_document(doc, output_path):
    try:
        doc.save(output_path)
        logging.info(f"Document saved successfully: {output_path}")
        print(f"Document saved successfully: {output_path}")
    except Exception as e:
        logging.error(f"Error saving document: {e}")
        print(f"Error saving document: {e}")

#main function to consolidate everything
def main():
    #Parse arguments
    parser = argparse.ArgumentParser(description="Retrieve Installed Drivers Information")
    parser.add_argument(
        "--output",
        help="Output file",
        default=os.getenv('DRIVERS_FILE', os.path.join(os.path.expanduser('~'), 'Documents', 'List_of_installed_drivers.docx'))
    )
    args = parser.parse_args()

    #Ensure output directory exists
    output_dir = os.path.dirname(args.output)
    os.makedirs(output_dir, exist_ok=True)

    #Set up logging
    setup_logging(output_dir)

    #Create a Word document
    doc = create_word_document()

    #Desired driver attributes
    desired_attributes = [
        "DeviceName", "DriverVersion", "FriendlyName", "InfName", "InstallDate",
        "IsSigned", "Location", "Manufacturer", "Name", "PDO", "DriverProviderName",
        "Signer", "Started", "StartMode", "Status", "SystemCreationClassName"
    ]

    #Get installed drivers
    drivers = get_installed_drivers()

    #Create a table in the document
    table = create_table(doc, desired_attributes)

    #Populate the table with driver data
    populate_table_with_driver_data(table, drivers, desired_attributes)

    #Save the document
    save_document(doc, args.output)


if __name__ == "__main__":
    main()
