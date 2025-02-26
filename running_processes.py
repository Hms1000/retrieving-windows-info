import os
import wmi
import sys
import logging
import argparse
from docx import Document

#configure logging
def configure_logging(output_dir):
    log_path = os.path.join(output_dir, 'List_of_running_processes.log')

    #ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    logging.basicConfig(
        filename=log_path,
        level=logging.INFO,
        format='%(asctime)s-%(levelname)s-%(message)s'
    )
    logging.info(f'Directory created successfully: {output_dir}')

#create a word document
def create_word_document():
    doc = Document()
    doc.add_heading('List of Running Processes', 0)
    return doc

#connect to WMI and retrieve running processes
def retrieve_running_processes():
    try:
        conn = wmi.WMI()
        processes = conn.Win32_Process()

        #check if processes are found
        if not processes:
            logging.warning('No processes found')
            print('No processes found')
            sys.exit(1)

        logging.info(f'Found {len(processes)} running processes')
        return processes

    except wmi.x_wmi as e:
        logging.error(f'Error running WMI processes: {e}')
        print(f'Error running WMI processes: {e}')
        sys.exit(1)
    except Exception as e:
        logging.error(f'Unexpected error occurred: {e}')
        print(f'Unexpected error occurred: {e}')
        sys.exit(1)

#create a table
def create_table(doc, desired_attributes):
    table = doc.add_table(rows=1, cols=len(desired_attributes))
    table.style = 'Table Grid'

    #add headers
    header_cells = table.rows[0].cells
    for index, attribute in enumerate(desired_attributes):
        header_cells[index].text = attribute

    return table

#populate the table with running processes information
def populate_table_with_running_process_information(table, processes, desired_attributes):
    for process in processes:
        row_cells = table.add_row().cells
        for index, attribute in enumerate(desired_attributes):
            try:
                value = getattr(process, attribute, None)

                #we handle empty and  missing values
                if value is None or (isinstance(value, str) and not value.strip()):
                    row_cells[index].text = 'N/A'
                else:
                    row_cells[index].text = str(value)

            except AttributeError as e:
                row_cells[index].text = 'Error'
                logging.error(f'Attribute Error: {e}')
                print(f'Attribute Error: {e}')

            except Exception as e:
                row_cells[index].text = 'Error'
                logging.error(f'Unexpected error occurred: {e}')
                print(f'Unexpected error occurred: {e}')

#save the wod document
def save_word_document(doc, output_path):
    try:
        doc.save(output_path)
        logging.info(f'Document saved successfully: {output_path}')
        print(f'Document saved successfully: {output_path}')
    except Exception as e:
        logging.error(f'Error saving document {output_path}: {e}')
        print(f'Error saving document {output_path}: {e}')

#main function to consolidate everything
def main():
    try:
        parser = argparse.ArgumentParser(description='Retrieve Running Processes')
        parser.add_argument(
            '--output', help='Output file', default=os.getenv('PROCESSES_FILE', os.path.join(os.path.expanduser('~'), 'Documents', 'List_of_running_processes.docx'))

        )

        args = parser.parse_args()
        print(f"Output path: {args.output}")

        #ensure the directory path exists
        output_dir = os.path.dirname(args.output)

        #configure logging
        configure_logging(output_dir)

        #create document
        doc = create_word_document()

        #retrieve running processes
        processes = retrieve_running_processes()

        #list of desired attributes
        desired_attributes = ['ProcessID', 'HandleCount', 'Name', 'ParentProcessId', 'ExecutablePath']

        #creating the table
        table = create_table(doc,desired_attributes)

        #populating the table
        populate_table_with_running_process_information(table, processes, desired_attributes)

        #saving the document
        save_word_document(doc, args.output)

    except Exception as e:
        logging.error(f'Error in main function: {e}')
        print(f'Error in main function: {e}')

if __name__ == '__main__':
    main()


